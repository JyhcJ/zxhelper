
import os
import requests
from bs4 import BeautifulSoup
from docx import Document

# 基础 URL
base_url = "https://tougao.12371.cn/tuijian.php"

# 请求头（模拟浏览器）
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
}

# 桌面路径
desktop_path = os.path.join(os.path.expanduser("~"), "Desktop/乡村振兴") #定义桌面上创建乡村振兴

# 创建文件夹
if not os.path.exists(desktop_path):  # 检查文件夹是否已存在
    os.makedirs(desktop_path)  # 创建文件夹（包括父目录）
    print(f"文件夹已创建: {desktop_path}")
else:
    print(f"文件夹已存在: {desktop_path}")
# 创建一个 Word 文档




# 遍历页面（从第 1 页到第 10 页）
for page in range(1, 10):
    print(f"正在抓取第 {page} 页...")
    url = f"{base_url}?&page={page}"  # 分页 URL
    response = requests.get(url, headers=headers) #读取网页数据
    response.encoding = 'utf-8' #设置网页编码格式

    # 解析网页内容
    soup = BeautifulSoup(response.text, 'html.parser')  #优化网页数据
    articles = soup.find_all('li', class_='bor_bom_none')  # 根据实际网页结构调整选择器

    # 如果没有找到稿件，说明已经抓取完所有页面
    if not articles:
        print("所有页面抓取完成！")
        break

    # 提取每篇稿件的信息
    for article in articles:
           # 找到标题和链接
        title_tag = article.find('span', class_='new_laigao').find('a') #寻找class为'new_laigao'下的span数据
        if title_tag:  # 确保找到标题和链接
            title = title_tag.text.strip()  # 标题
            link = title_tag['href']  # 链接
        
        # 过滤出标题中包含“乡村振兴”的稿件

        if "乡村振兴" in title:
            print(f"找到相关稿件: {title}")

            # 将标题和链接添加到 Word 文档
            #doc.add_paragraph(f"{title}  : https://tougao.12371.cn/{link}")
       
            print(f"正在抓取 {title} 正文...")
            url1 = "https://tougao.12371.cn/"+link      #设置有关乡村振兴的链接
            #textUrl = f"&{url1}"  # 分页 URL
            response = requests.get(url1, headers=headers)   #打开有关乡村振兴的链接
            response.encoding = 'utf-8'                      #设置网页编码格式

            soup1 = BeautifulSoup(response.text, 'html.parser') #优化网页数据
            word_div = soup1.find('div', class_='word')  # 根据实际网页结构调整选择器
            if word_div:
                text_content = word_div.get_text(separator='\n', strip=True)  # 使用换行符分隔文本
                #text_content = word_div
           
            doc = Document() #定义doc对象
            doc.add_heading('乡村振兴相关稿件链接:'+url1, level=1)  # doc添加标题
         
            doc.add_paragraph(f"{text_content}")            # doc添加内容
         
            output_file = os.path.join(desktop_path, title+".docx") # doc创建文件
            doc.save(output_file)                           #doc保存数据

print(f"所有链接已保存到: {output_file}")